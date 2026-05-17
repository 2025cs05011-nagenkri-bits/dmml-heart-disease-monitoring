"""
generate_evidently_demo_report.py

Builds the Tool 2: Evidently AI demo report in the same style as the
team's Tool 1: Great Expectations report (Great_Expectations_Demo_
Report.docx). Skips the group/members header so the file can be
appended manually onto the combined submission. Output:

  reports/Evidently_Demo_Report.docx          (~5-6 page demo report)

Reuses the architecture diagram and screenshots already produced for
the longer DMML_Evidently_Report.docx.

Usage (from DMML_Evidently_Demo/):
  python -m scripts.generate_evidently_demo_report

Requires: python-docx==1.1.2
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = ROOT / "reports" / "figures"
SHOT_DIR = ROOT / "reports" / "screenshots" / "evidently"
OUT_DOCX = ROOT / "reports" / "Evidently_Demo_Report.docx"
OVERVIEW_PNG = FIG_DIR / "evidently_overview.png"


# ---------------------- Conceptual overview diagram ----------------------

def draw_evidently_overview(out_path: Path) -> None:
    """Generic Evidently flow: two datasets -> engine -> three outputs.

    Tool-agnostic and pipeline-agnostic; appears in the Overview
    section to give readers a one-glance mental model of what
    Evidently does before the project-specific architecture diagram.
    """
    fig, ax = plt.subplots(figsize=(13, 6.5))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 7)
    ax.set_aspect("equal")
    ax.axis("off")

    # Inputs (left)
    inputs = [
        (0.4, 4.5, 2.6, 1.2,
         "Reference dataset\n(training-time snapshot)", "#bbdefb"),
        (0.4, 1.3, 2.6, 1.2,
         "Current dataset\n(production / today's data)", "#e3f2fd"),
    ]
    for x, y, w, h, label, color in inputs:
        ax.add_patch(mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.04",
            facecolor=color, edgecolor="#3b3b3b",
            linewidth=1.2, zorder=2,
        ))
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center",
                fontsize=10, zorder=3)

    # Engine (middle) - large box with three preset rows
    ex, ey, ew, eh = 4.0, 1.0, 4.6, 5.2
    ax.add_patch(mpatches.FancyBboxPatch(
        (ex, ey), ew, eh, boxstyle="round,pad=0.08",
        facecolor="#e1bee7", edgecolor="#3b3b3b",
        linewidth=1.4, zorder=2,
    ))
    ax.text(ex + ew / 2, ey + eh - 0.45, "Evidently AI engine",
            ha="center", va="center", fontsize=12, fontweight="bold",
            zorder=3)
    ax.text(ex + ew / 2, ey + eh - 0.95,
            "evidently.report.Report(metrics=[...])",
            ha="center", va="center", fontsize=9, style="italic",
            color="#444", zorder=3)

    # Three preset rows inside engine
    presets = [
        (eh - 1.7, "DataDriftPreset",
         "per-column distribution tests"),
        (eh - 2.7, "TargetDriftPreset",
         "label / target prevalence shift"),
        (eh - 3.7, "DataQualityPreset",
         "missing values, ranges, summaries"),
    ]
    for dy, name, desc in presets:
        py = ey + dy
        ax.add_patch(mpatches.FancyBboxPatch(
            (ex + 0.25, py), ew - 0.5, 0.85,
            boxstyle="round,pad=0.03",
            facecolor="#f3e5f5", edgecolor="#7b1fa2",
            linewidth=0.9, zorder=3,
        ))
        ax.text(ex + 0.4, py + 0.6, name, ha="left", va="center",
                fontsize=9.5, fontweight="bold", zorder=4)
        ax.text(ex + 0.4, py + 0.25, desc, ha="left", va="center",
                fontsize=8.5, color="#555", zorder=4)

    # Statistical tests footnote inside engine
    ax.text(ex + ew / 2, ey + 0.35,
            "tests: Wasserstein  |  Chi-square  |  K-S  |  PSI",
            ha="center", va="center", fontsize=8.5,
            color="#444", style="italic", zorder=3)

    # Outputs (right)
    outputs = [
        (9.6, 5.2, 3.0, 1.1, "HTML report",
         "human-readable, interactive", "#fff3e0"),
        (9.6, 3.1, 3.0, 1.1, "JSON metrics",
         "machine-readable for CI gates", "#fff3e0"),
        (9.6, 1.0, 3.0, 1.1, "dataset_drift verdict",
         "True / False + share of\ndrifted columns", "#c8e6c9"),
    ]
    for x, y, w, h, title, sub, color in outputs:
        ax.add_patch(mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.04",
            facecolor=color, edgecolor="#3b3b3b",
            linewidth=1.2, zorder=2,
        ))
        ax.text(x + w / 2, y + h - 0.32, title, ha="center", va="center",
                fontsize=10.5, fontweight="bold", zorder=3)
        ax.text(x + w / 2, y + 0.32, sub, ha="center", va="center",
                fontsize=8.5, color="#555", zorder=3)

    arrow_kw = dict(arrowstyle="-|>", color="#555", lw=1.2,
                    alpha=0.9, mutation_scale=14)
    # Inputs -> engine
    ax.annotate("", xy=(ex, ey + eh - 1.3), xytext=(3.0, 5.1),
                arrowprops=arrow_kw, zorder=2)
    ax.annotate("", xy=(ex, ey + 1.3), xytext=(3.0, 1.9),
                arrowprops=arrow_kw, zorder=2)
    # Engine -> outputs
    ax.annotate("", xy=(9.6, 5.75), xytext=(ex + ew, ey + eh - 1.3),
                arrowprops=arrow_kw, zorder=2)
    ax.annotate("", xy=(9.6, 3.65), xytext=(ex + ew, ey + eh / 2),
                arrowprops=arrow_kw, zorder=2)
    ax.annotate("", xy=(9.6, 1.55), xytext=(ex + ew, ey + 1.3),
                arrowprops=arrow_kw, zorder=2)

    ax.set_title("Evidently AI - what the library does",
                 fontsize=13, fontweight="bold", pad=10)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


# ------------------------------ docx helpers ------------------------------

def _set_run(run, *, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    return h


def add_para(doc, text, *, size=11, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    _set_run(r, size=size, bold=bold, italic=italic)
    return p


def add_bold_inline(doc, lead_bold, tail):
    """Paragraph that starts with a bold lead and continues in normal."""
    p = doc.add_paragraph()
    r1 = p.add_run(lead_bold)
    _set_run(r1, size=11, bold=True)
    r2 = p.add_run(tail)
    _set_run(r2, size=11)
    return p


def add_bullets(doc, items):
    """items can be plain strings, or (bold_lead, tail) tuples."""
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            lead, tail = item
            r1 = p.add_run(lead)
            _set_run(r1, size=11, bold=True)
            r2 = p.add_run(tail)
            _set_run(r2, size=11)
        else:
            r = p.add_run(item)
            _set_run(r, size=11)


def add_figure(doc, img_path: Path, caption: str, width_cm=14.5):
    if not img_path.exists():
        return
    doc.add_picture(str(img_path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    _set_run(r, size=10, italic=True, color=(0x55, 0x55, 0x55))


def add_steps_table(doc, header, rows):
    """3-column table styled to match the GE report's step tables."""
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for cell, label in zip(table.rows[0].cells, header):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(label)
        _set_run(r, size=11, bold=True)

    for ri, row in enumerate(rows, start=1):
        for cell, value in zip(table.rows[ri].cells, row):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            r = p.add_run(value)
            _set_run(r, size=10.5)
    return table



# ------------------------------ Report body ------------------------------

def build_report() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Tool 2 banner ----
    add_para(doc, "Tool 2: Evidently AI", size=16, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---- Overview ----
    add_heading(doc, "Overview", level=1)
    add_para(doc,
        "Once a machine-learning model is shipped, the world keeps "
        "moving - sensors get recalibrated, customer behaviour shifts, "
        "upstream pipelines change - and the production input "
        "distribution slowly drifts away from the one the model was "
        "trained on. Without an automated guard, accuracy degrades "
        "silently while the unit-test suite continues to pass, because "
        "tests verify code behaviour, not data behaviour. This is where "
        "data drift detection plays an important role.")
    add_para(doc,
        "Evidently AI is an open-source Python library for ML "
        "observability that has become popular for catching distribution "
        "shifts before they reach production.")
    add_para(doc,
        "Its Python-based framework is designed to help data teams "
        "monitor the health of their data and models. Users describe a "
        "reference dataset (the distribution the model was trained on) "
        "and a current dataset (today's inputs); Evidently then runs a "
        "battery of statistical tests, produces an HTML report for "
        "humans and a JSON summary for machines, and reports a single "
        "dataset-level drift verdict that a CI gate can act on.")
    add_para(doc, "Some benefits of Evidently AI include:")
    add_bullets(doc, [
        ("Pre-built drift presets - ",
         "DataDriftPreset (per-column distribution tests), "
         "TargetDriftPreset (label-prevalence shifts) and "
         "DataQualityPreset (descriptive profiling) work out of the "
         "box without writing custom test logic."),
        ("Per-column statistical tests - ",
         "Wasserstein distance for numeric features, Chi-square for "
         "categoricals, with configurable thresholds and a "
         "dataset-level share-of-drifted-columns verdict."),
        ("Dual-format output - ",
         "Reports as HTML for human inspection and JSON for CI gates, "
         "both written from the same Report object."),
        ("Pipeline integration - ",
         "Pure-Python API with no external services; trivial to wire "
         "into pytest, GitHub Actions, Airflow or any CI runner."),
        ("Frozen-reference pattern - ",
         "Designed around comparing a current dataset to a stable "
         "reference, which maps cleanly onto a version-controlled "
         "training-snapshot workflow."),
        ("Open source and self-hosted - ",
         "Apache 2.0, no vendor lock-in, runs entirely inside the "
         "project's existing Python environment."),
    ])
    add_figure(doc, OVERVIEW_PNG,
               "Figure - At a glance: Evidently compares a reference "
               "and a current dataset using its built-in presets and "
               "statistical tests, then emits an HTML report, a JSON "
               "metrics file and a single dataset_drift verdict that "
               "downstream pipelines can act on.", width_cm=15.5)
    add_para(doc,
        "More details on the Evidently tool: "
        "https://docs.evidentlyai.com/")
    add_para(doc, "******************************",
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc,
        "Now, let's look at an end-to-end drift detection on the Heart "
        "Disease UCI dataset using the Evidently AI tool!", italic=True)
    add_para(doc,
        "Source: scripts/run_evidently.py and src/monitoring/drift.py "
        "in the DMML_Evidently_Demo repository.", italic=True)

    # ---- 1. Introduction ----
    add_heading(doc, "1. Introduction", level=1)
    add_para(doc,
        "This report summarises the Evidently AI integration into the "
        "Heart Disease UCI MLOps pipeline, a hands-on walkthrough of "
        "data profiling and drift detection wired into a real CI "
        "workflow. The integration extends an existing pipeline "
        "(scikit-learn training with MLflow tracking, Flask serving, "
        "Helm-deployed Kubernetes target, GitHub Actions CI) with a "
        "new drift gate that compares every CI run's training data "
        "against a frozen reference snapshot. If the comparison flags "
        "a statistically significant divergence, the build fails "
        "before a stale model can be packaged or deployed.")
    add_figure(doc, FIG_DIR / "dmml_architecture.png",
               "Figure 1 - End-to-end Evidently AI drift gate "
               "architecture: data sources, engine, CI gating and "
               "outputs.", width_cm=15.5)

    # ---- 2. Environment Setup and Reference Snapshot ----
    add_heading(doc,
        "2. Environment Setup and Reference Snapshot (Steps 1-3)",
        level=1)
    add_para(doc,
        "The integration begins by adding evidently to requirements.txt "
        "and creating a frozen baseline. The dataset is loaded with "
        "pandas; column types are declared explicitly via Evidently's "
        "ColumnMapping, sourced from the same constants used by the "
        "production preprocessing pipeline (src/data/preprocess.py). "
        "This guarantees that the drift detector and the production "
        "preprocessor always agree on which features are numeric, "
        "categorical and binary, eliminating a common source of "
        "false-positive drift caused by misclassified columns.")
    add_steps_table(doc,
        ["Step", "Operation", "Purpose"],
        [
            ("Installation",
             "pip install evidently==0.4.40",
             "Installs the drift-detection library"),
            ("Import library",
             "from evidently.report import Report",
             "Loads the Report API and the drift / quality presets"),
            ("Load reference",
             "pd.read_csv('data/reference/"
             "reference_train_v1.0.0.csv')",
             "Loads the frozen training snapshot from git"),
            ("Freeze snapshot",
             "python scripts/freeze_reference.py --version v1.0.0",
             "Pins train.csv plus SHA-256 metadata to git as the "
             "immutable baseline"),
        ])


    # ---- 3. Core Workflow ----
    add_heading(doc, "3. Core Workflow (Steps 4-10)", level=1)
    add_para(doc,
        "The drift module (src/monitoring/drift.py) exposes three pure "
        "functions that wrap Evidently presets: build_data_profile, "
        "build_drift_report and extract_drift_summary. Reports are "
        "written to stable filenames (profile.html / profile.json and "
        "drift.html / drift.json) and overwritten each run, so "
        "reports/evidently/ always contains exactly five files. "
        "scripts/run_evidently.py is the CLI entry point that invokes "
        "the engine, prints the headline verdict and exits with code 1 "
        "when drift is detected (unless --no-fail-on-drift is set).")
    add_steps_table(doc,
        ["Step", "Check", "Key API"],
        [
            ("4", "Build per-column profile of reference",
             "DataQualityPreset()"),
            ("5", "Compare current vs reference distributions",
             "DataDriftPreset()"),
            ("6", "Detect target prevalence shift",
             "TargetDriftPreset()"),
            ("7", "Persist HTML and JSON reports",
             "report.save_html(...) / save_json(...)"),
            ("8", "Parse JSON for the CI verdict",
             "extract_drift_summary(json_path)"),
            ("9", "Pipeline gate (fail-fast on drift)",
             "sys.exit(1) if dataset_drift else sys.exit(0)"),
            ("10", "Visual report attached to CI artifact",
             "actions/upload-artifact, if: always()"),
        ])
    add_figure(doc, SHOT_DIR / "13_profile_html.png",
               "Figure 2 - DataQualityPreset profile.html: per-column "
               "statistics, missing values and distribution summaries "
               "for the reference set (Step 4).", width_cm=14.5)

    # ---- 4. Drift Demonstration & CI Failure Injection ----
    add_heading(doc,
        "4. Drift Demonstration and CI Failure Injection", level=1)
    add_para(doc,
        "To prove the gate actually rejects bad data, "
        "scripts/generate_drift_demo.py fabricates a current dataset "
        "with deliberately injected drift. Numeric features are shifted "
        "by 1.5 standard deviations, around 30 percent of categorical "
        "entries are flipped to other valid levels, and the target "
        "prevalence is skewed downward. The CI workflow exposes a "
        "manual workflow_dispatch toggle that runs the gate against "
        "this dirty data on demand.")
    add_para(doc,
        "Running run_evidently.py against the synthetic input yields "
        "dataset_drift = True with share_of_drifted_columns = 0.7143 "
        "(10 of 14 columns flagged). The script exits with code 1, "
        "GitHub marks the step failed, and the dependent Docker build "
        "job is skipped. Despite the failure, the artifact upload step "
        "still runs (if: always()) so the HTML reports remain "
        "downloadable from the failed run page.")
    add_para(doc, "Step 12 - injected drift and the detectors that trip:")
    add_steps_table(doc,
        ["#", "Injected drift", "Detector that trips"],
        [
            ("1", "Numeric features shifted by +1.5 sigma "
             "(age, chol, thalach, trestbps, oldpeak)",
             "DataDriftPreset (Wasserstein distance)"),
            ("2", "~30% categorical flips on cp, restecg, slope, thal",
             "DataDriftPreset (Chi-square test)"),
            ("3", "Target prevalence skewed by ~10 pp",
             "TargetDriftPreset"),
            ("4", "Combined effect across 14 columns",
             "dataset_drift = True, share = 0.7143"),
            ("5", "Exit code 1 propagates through CI",
             "Docker build job is skipped"),
        ])
    add_figure(doc, SHOT_DIR / "09_failed_step_log.png",
               "Figure 3 - Failed CI run on the synthetic drifted "
               "input: dataset_drift = True, share = 0.7143, exit code "
               "1, build marked red.", width_cm=14.5)

    # ---- 5. Advanced Features ----
    add_heading(doc,
        "5. Advanced Features (Steps 13-17)", level=1)
    add_para(doc,
        "The second half of the integration demonstrates the patterns "
        "that distinguish a production-ready drift gate from a one-off "
        "script: a versioned reference snapshot, a three-mode CI gating "
        "strategy, unconditional artifact upload, per-column "
        "drill-down and stable filenames.")
    add_steps_table(doc,
        ["Step", "Capability", "Implementation"],
        [
            ("13", "Frozen reference with SHA-256 metadata",
             "scripts/freeze_reference.py + "
             "reference_train_<version>.metadata.json"),
            ("14", "Three-mode CI gate (real / aux demo / failure demo)",
             "workflow_dispatch input with truthy if: condition"),
            ("15", "Unconditional artifact upload (works on red builds)",
             "if: always() on actions/upload-artifact@v4"),
            ("16", "Per-column distribution overlays",
             "drift.html per-column drill-down section"),
            ("17", "Stable filenames overwritten each run",
             "save_report() with fixed paths under reports/evidently/"),
        ])
    add_para(doc,
        "Step 13 highlight - frozen reference. Each baseline is "
        "committed as reference_train_<version>.csv plus a JSON sidecar "
        "capturing version, sha256, n_rows, target_prevalence, full "
        "schema and the git HEAD at freeze time. The current baseline "
        "v1.0.0 covers 242 rows x 14 columns at SHA-256 prefix "
        "0e538464. Bumping the baseline is one freeze command plus one "
        "CI line change, and the metadata file becomes the audit trail "
        "of why.")
    add_para(doc, "Selected Step 14 finding:")
    add_bullets(doc, [
        "workflow_dispatch boolean inputs are sometimes stringified by "
        "GitHub, so the conditional uses a truthy check rather than a "
        "literal == true comparison; otherwise the failure-demo step "
        "silently skips even when the toggle is on.",
        "All three CI modes share the same Evidently engine; only the "
        "input dataset and the --no-fail-on-drift flag differ. The "
        "real gate runs on every push, the auxiliary demo never blocks, "
        "and the failure demo blocks only when the manual toggle is on.",
    ])
    add_para(doc,
        "Step 15 - artifact upload. Even when the drift gate fails the "
        "build, actions/upload-artifact@v4 still executes because of "
        "if: always(). Engineers can download evidently-reports-py3.9 "
        "from the red run page and inspect drift.html without "
        "re-running anything locally.")
    add_figure(doc, SHOT_DIR / "12_drift_html_per_column.png",
               "Figure 4 - drift.html per-column drill-down: reference "
               "vs current distribution overlay for a drifted feature, "
               "with the Wasserstein test result and verdict.",
               width_cm=14.5)
    add_para(doc, "Deliverables produced by the integration:")
    add_bullets(doc, [
        "src/monitoring/drift.py - engine module wrapping the three "
        "Evidently presets behind pure functions.",
        "scripts/run_evidently.py - CLI entry point with "
        "--no-fail-on-drift flag for non-blocking runs.",
        "scripts/freeze_reference.py - versions a training snapshot to "
        "git with SHA-256 audit metadata.",
        "data/reference/reference_train_v1.0.0.csv plus "
        ".metadata.json - the frozen baseline used by every CI gate.",
        "reports/evidently/{profile,drift}.{html,json} plus "
        "summary.json - five stable files overwritten on each run.",
    ])


    # ---- 6. Conclusion ----
    add_heading(doc, "6. Conclusion", level=1)
    add_para(doc,
        "The Evidently AI integration provides a complete walkthrough "
        "of automated data-drift detection wired into a real CI "
        "pipeline. It demonstrates how a frozen training snapshot, a "
        "small Python engine and three GitHub Actions steps are enough "
        "to turn drift detection from an offline analytics task into a "
        "build-time gate that physically prevents a stale model from "
        "shipping. The same engine powers a one-click stakeholder demo "
        "of both the healthy path and the rejection path on the same "
        "commit, and the Evidently HTML reports remain inspectable on "
        "every run regardless of pass or fail.")
    add_para(doc,
        "Overall, the implementation shows how Evidently AI can detect "
        "distributional drift before a stale model is deployed, give "
        "engineers a downloadable HTML report on every run regardless "
        "of build outcome, and complement code-level tests with a "
        "data-level test that catches the failure modes unit tests "
        "cannot see.")
    return doc


# ------------------------------ entry point ------------------------------

def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    print(f"[1/2] writing overview diagram -> {OVERVIEW_PNG}")
    draw_evidently_overview(OVERVIEW_PNG)
    print(f"[2/2] writing demo report -> {OUT_DOCX}")
    doc = build_report()
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    size_kb = OUT_DOCX.stat().st_size / 1024
    print(f"        wrote {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
