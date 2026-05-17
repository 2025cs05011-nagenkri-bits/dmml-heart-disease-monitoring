"""
generate_dmml_report.py

Builds a focused ~7-page DMML / Evidently AI report covering data
profiling and drift detection on top of the existing Heart Disease
MLOps pipeline. Output:

  reports/figures/dmml_architecture.png       (matplotlib block diagram)
  reports/DMML_Evidently_Report.docx          (~7-page final report)

Screenshots are pulled from reports/screenshots/evidently/ (see the
README in that folder for filenames and capture guidance). Missing
screenshots are silently skipped so the report can be rebuilt
incrementally as captures are added.

Usage (from DMML_Evidently_Demo/):
  python -m scripts.generate_dmml_report

Requires: python-docx==1.1.2  (in addition to requirements-dev.txt)
"""
from __future__ import annotations

from pathlib import Path

import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIG_DIR = ROOT / "reports" / "figures"
SHOT_DIR = ROOT / "reports" / "screenshots" / "evidently"
OUT_DOCX = ROOT / "reports" / "DMML_Evidently_Report.docx"

# Group members (Student Name, BITS ID). Listed in submission order.
GROUP_MEMBERS = [
    ("Bhaskar Borah", "2025CS05039"),
    ("Ramya K S", "2025CS05055"),
    ("Nikhil Rahul Ekka", "2025CS05012"),
    ("Nagendren K", "2025CS05011"),
]

# Screenshot manifest: (filename, caption_tail). Numbering is assigned
# at render time so the figure counter stays continuous across sections.
SHOTS_REFERENCE = [
    ("01_metadata_json.png",
     "Frozen reference metadata file - schema, SHA-256, target prevalence "
     "and git provenance committed alongside the snapshot."),
]
SHOTS_ENGINE = [
    ("13_profile_html.png",
     "DataQualityPreset profile HTML - per-column statistics, missing "
     "values and distribution summaries on the reference set."),
    ("14_evidently_folder.png",
     "reports/evidently/ directory after a run - exactly five stable files "
     "(profile + drift HTML/JSON pairs and the summary), overwritten in "
     "place each invocation."),
]
SHOTS_CI = [
    ("04_workflow_dispatch_dropdown.png",
     "GitHub Actions workflow_dispatch panel - the 'demo_fail_on_drift' "
     "checkbox toggles between the healthy and failure scenarios."),
    ("08_debug_inputs_step.png",
     "Debug step output confirming the workflow_dispatch input value that "
     "the conditional 'if' clauses gate on."),
]
SHOTS_HEALTHY = [
    ("02_local_clean_run.png",
     "Local invocation of run_evidently.py against the frozen reference - "
     "dataset_drift=False, share_of_drifted_columns=0.0000, exit code 0."),
    ("06_green_run_steps.png",
     "Successful CI run - the real drift gate step executes against the "
     "frozen reference and reports zero drift."),
    ("13_1_Drift_html_No_drift.png",
     "drift.html for the healthy path - dataset-level verdict reads 'no "
     "drift' with zero drifted columns when current data matches the "
     "frozen reference."),
]
SHOTS_DRIFTED = [
    ("03_local_drifted_run.png",
     "Local invocation against the synthetic drifted current dataset - "
     "dataset_drift=True, exit code 1; identical behaviour to the CI gate."),
    ("07_red_run_overview.png",
     "Failed CI run summary triggered via workflow_dispatch with the "
     "demo_fail_on_drift toggle on."),
    ("09_failed_step_log.png",
     "Expanded 'Drift gate FAILURE demo' step log - dataset_drift=True, "
     "share=0.7143, FAIL message and non-zero exit code."),
]
SHOTS_OUTPUTS = [
    ("05_runs_list_green_red.png",
     "Workflow runs list - green push run and red workflow_dispatch run on "
     "the same commit, illustrating the gate's binary verdict."),
    ("10_artifacts_section.png",
     "Artifacts panel of the failed run - the Evidently report bundle is "
     "still uploaded thanks to 'if: always()' on the upload step."),
    ("11_drift_html_overview.png",
     "drift.html overview - dataset-level verdict, drifted column count "
     "and share-of-drifted-columns at the top of the report."),
    ("12_drift_html_per_column.png",
     "drift.html per-column section - reference vs current distribution "
     "overlay for a drifted feature with the statistical test result."),
]


# ----------------------- Architecture diagram -----------------------

ARCH_BOXES = [
    # (x, y, w, h, label, color)
    # ---- Band 1: Data Sources & Prep (y=8.15, h=0.9) ----
    (0.30, 8.15, 1.85, 0.9, "Raw CSV\nCleveland\n303x14", "#e3f2fd"),
    (2.40, 8.15, 1.85, 0.9, "prepare.py\nseed=42\nstratified", "#fff3e0"),
    (4.50, 8.15, 1.85, 0.9, "train.csv\n242x14", "#fff3e0"),
    (6.85, 8.15, 2.55, 0.9, "Frozen Reference\nv1.0.0 (in git)\nSHA: 0e538464", "#bbdefb"),
    (9.65, 8.15, 3.05, 0.9, "Synthetic Drift\ngenerate_drift_demo.py", "#e3f2fd"),

    # ---- Band 2: Evidently Engine (y=5.9, h=1.05) ----
    (0.30, 5.90, 3.85, 1.05, "build_data_profile()\nDataQualityPreset", "#f3e5f5"),
    (4.30, 5.90, 4.10, 1.05, "build_drift_report()\nDataDriftPreset +\nTargetDriftPreset", "#e1bee7"),
    (8.55, 5.90, 4.15, 1.05, "extract_drift_summary()\nWasserstein + Chi-square", "#f3e5f5"),

    # ---- Band 3: GitHub Actions CI (y=3.65, h=1.05) ----
    (0.30, 3.65, 3.85, 1.05, "Real Gate\nevery push / PR\nfrozen ref vs current", "#c8e6c9"),
    (4.30, 3.65, 4.10, 1.05, "Auxiliary Demo\nmanual, toggle OFF\n(--no-fail-on-drift)", "#c8e6c9"),
    (8.55, 3.65, 4.15, 1.05, "FAILURE Demo\nmanual, toggle ON\nexits non-zero on drift", "#a5d6a7"),

    # ---- Band 4: Outputs (y=1.40, h=1.0) ----
    (0.30, 1.40, 3.85, 1.0, "Reports (5 files)\n+ exit 0 -> PASS\nDocker proceeds", "#c8e6c9"),
    (4.30, 1.40, 4.10, 1.0, "Reports (5 files)\nalways exit 0\nartifact uploaded", "#dcedc8"),
    (8.55, 1.40, 4.15, 1.0, "Reports (5 files)\n+ exit 1 -> FAIL\nDocker skipped", "#ffcdd2"),
]


# Each engine box is horizontally aligned with one CI runner, which is
# horizontally aligned with one outputs box. This lets every cross-band
# arrow be a clean vertical line and avoids diagonal crisscrossing.
ARCH_ARROWS = [
    # ---- Band 1 internal (left-to-right pipeline) ----
    (2.15, 8.60, 2.40, 8.60),       # Raw CSV  -> prepare.py
    (4.25, 8.60, 4.50, 8.60),       # prepare  -> train.csv

    # ---- Band 1 -> Band 2 (entering box tops at y=6.95) ----
    (5.42, 8.15, 2.22, 6.95),       # train.csv -> build_data_profile
    (5.42, 8.15, 6.35, 6.95),       # train.csv -> build_drift_report
    (8.12, 8.15, 6.35, 6.95),       # Frozen Reference -> build_drift_report
    (11.17, 8.15, 6.35, 6.95),      # Synthetic Drift -> build_drift_report

    # ---- Band 2 internal (drift report -> summary parser) ----
    (8.40, 6.42, 8.55, 6.42),       # build_drift_report -> extract_drift_summary

    # ---- Band 2 -> Band 3 (clean vertical arrows) ----
    (2.22, 5.90, 2.22, 4.70),       # build_data_profile -> Real Gate
    (6.35, 5.90, 6.35, 4.70),       # build_drift_report -> Auxiliary Demo
    (10.62, 5.90, 10.62, 4.70),     # extract_drift_summary -> Failure Demo

    # ---- Band 3 -> Band 4 (clean vertical arrows) ----
    (2.22, 3.65, 2.22, 2.40),       # Real Gate -> Reports + PASS
    (6.35, 3.65, 6.35, 2.40),       # Auxiliary Demo -> Reports (always pass)
    (10.62, 3.65, 10.62, 2.40),     # Failure Demo -> Reports + FAIL
]

BANDS = [
    (7.95, 1.95, "Data Sources & Prep"),
    (5.70, 1.95, "Evidently Engine - src/monitoring/drift.py"),
    (3.45, 1.95, "GitHub Actions CI - .github/workflows/ci.yml"),
    (1.20, 1.95, "Outputs"),
]


def draw_architecture(out_path: Path) -> None:
    fig, ax = plt.subplots(figsize=(14, 10))
    ax.set_xlim(0, 13)
    ax.set_ylim(0.5, 10.2)
    ax.set_aspect("equal")
    ax.axis("off")

    # Bands - subtle background for each tier. Title sits at the top
    # edge with va='top' so the text never overlaps any inner box.
    for y, h, title in BANDS:
        ax.add_patch(mpatches.FancyBboxPatch(
            (0.15, y), 12.7, h, boxstyle="round,pad=0.02",
            facecolor="#f5f5f7", edgecolor="#d0d0d4",
            linewidth=1, zorder=0,
        ))
        ax.text(0.30, y + h - 0.10, title, fontsize=10.5,
                fontweight="bold", color="#444",
                ha="left", va="top", zorder=1)

    for x, y, w, h, label, color in ARCH_BOXES:
        ax.add_patch(mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.04",
            facecolor=color, edgecolor="#3b3b3b",
            linewidth=1.1, zorder=2,
        ))
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center",
                fontsize=8.8, zorder=3)

    for x1, y1, x2, y2 in ARCH_ARROWS:
        ax.annotate(
            "", xy=(x2, y2), xytext=(x1, y1),
            arrowprops=dict(arrowstyle="-|>", color="#555",
                            lw=1.1, alpha=0.85,
                            mutation_scale=12),
            zorder=2,
        )

    ax.set_title("Evidently AI Drift Gate - End-to-End Architecture",
                 fontsize=13.5, fontweight="bold", pad=12)
    fig.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)


# ------------------------------ docx helpers ------------------------------

def _set_run(run, *, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
    return h


def add_para(doc, text, *, size=11, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    _set_run(r, size=size, bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        _set_run(r, size=11)


def add_figure(doc, img_path: Path, caption: str, width_cm=15.0):
    if not img_path.exists():
        return
    doc.add_picture(str(img_path), width=Cm(width_cm))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    _set_run(r, size=9, italic=True, color=(0x55, 0x55, 0x55))


def add_group_members_table(doc, members):
    """Render a centered two-column table: Student Name | BITS ID."""
    table = doc.add_table(rows=1 + len(members), cols=2)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_cells = table.rows[0].cells
    for cell, label in zip(header_cells, ("Student Name", "BITS ID")):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(label)
        _set_run(r, size=11, bold=True)

    for row_idx, (name, bits_id) in enumerate(members, start=1):
        row_cells = table.rows[row_idx].cells
        for cell, value in zip(row_cells, (name, bits_id)):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            _set_run(r, size=11)
    return table


def add_screenshots(doc, shots, start_idx: int, width_cm: float = 14.5) -> int:
    idx = start_idx
    for fname, caption_tail in shots:
        path = SHOT_DIR / fname
        if not path.exists():
            continue
        add_figure(doc, path, f"Figure {idx}. {caption_tail}",
                   width_cm=width_cm)
        idx += 1
    return idx


# ------------------------------ Report body ------------------------------

def build_report() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Cover ----
    add_para(doc, "DMML Tooling Integration", size=14, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Data Profiling and Drift Detection with",
             size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Evidently AI on the Heart Disease MLOps Pipeline",
             size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_figure(doc, FIG_DIR / "dmml_architecture.png",
               "Figure 1. End-to-end Evidently AI drift gate architecture - "
               "data sources, engine, CI gating and outputs.", width_cm=16)
    doc.add_paragraph()
    add_para(doc, "Group Members Name with Student ID",
             size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_group_members_table(doc, GROUP_MEMBERS)
    doc.add_paragraph()
    add_para(doc, "Programme: M.Tech Software Engineering, BITS Pilani WILP",
             size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Repository: github.com/2025cs05011-nagenkri-bits/"
             "dmml-heart-disease-monitoring", size=10, italic=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ---- 1. Environment Recap ----
    add_heading(doc, "1. Environment and Project Recap", level=1)
    add_para(doc,
        "This activity extends the existing Heart Disease UCI MLOps "
        "pipeline (Assignment 1) with a Data Mining and Machine Learning "
        "(DMML) tool, Evidently AI, focused on data profiling and drift "
        "detection. The base pipeline already covers data preparation "
        "(src/data/prepare.py), training with cross-validation and MLflow "
        "tracking (src/training/train.py), a Flask serving layer with "
        "Prometheus metrics, a Helm-deployed Kubernetes target and a "
        "GitHub Actions CI workflow. No changes were required to the "
        "model, container or deployment surface; the integration is "
        "additive and isolated to a new monitoring sub-package, two "
        "helper scripts and three additional CI steps.")
    add_para(doc, "Tooling stack used in this activity:")
    add_bullets(doc, [
        "Evidently AI 0.4.40 - the DMML library providing DataDriftPreset, "
        "TargetDriftPreset and DataQualityPreset.",
        "Python 3.9 in a project-local venv (DMML_Evidently_Demo/venv); "
        "scikit-learn 1.6.1 and pandas 2.3.3 are inherited from the "
        "production runtime image and pinned via requirements.txt.",
        "GitHub Actions for continuous integration, with the workflow "
        "file at .github/workflows/ci.yml driving lint, unit tests, "
        "training reproduction and the new drift gate.",
        "Git as the canonical store of the immutable reference snapshot "
        "(data/reference/) used for every drift comparison.",
    ])
    add_para(doc,
        "All work is contained in the DMML_Evidently_Demo/ subproject of "
        "the DMML_Project workspace. Reproducing this report end-to-end "
        "from a clean checkout requires only `pip install -r "
        "requirements.txt -r requirements-dev.txt` followed by `python "
        "scripts/run_evidently.py` and `python -m scripts.generate_"
        "dmml_report`.")
    doc.add_page_break()

    # ---- 2. Why Drift Detection ----
    add_heading(doc, "2. Why Drift Detection", level=1)
    add_para(doc,
        "A model trained on a historical snapshot of the input "
        "distribution will silently lose accuracy when the production "
        "data distribution diverges from that snapshot. The divergence "
        "may originate upstream (a sensor recalibration, a label-mapping "
        "change, a new data provider), or it may be introduced by an "
        "engineering accident (a developer modifies the data preparation "
        "script, replaces the raw file or bumps the random seed). In "
        "either case, the deployed model's predictions become "
        "increasingly unreliable while the test suite continues to pass, "
        "because unit tests verify code behaviour, not data behaviour.")
    add_para(doc,
        "The standard mitigation is to treat data as a first-class build "
        "input: capture the distribution the production model was trained "
        "on as an immutable artefact, then on every change recompute the "
        "distribution and compare. If the comparison flags a statistically "
        "significant divergence, fail the build before the bad model can "
        "be packaged and deployed. This activity implements exactly that "
        "pattern, with Evidently AI providing the statistical machinery "
        "and the existing GitHub Actions workflow providing the gate.")
    add_heading(doc, "2.1 What Evidently AI provides", level=2)
    add_para(doc,
        "Evidently AI is an open-source Python library for ML "
        "observability. The features used in this activity are:")
    add_bullets(doc, [
        "DataDriftPreset - per-column statistical tests (Wasserstein "
        "distance for numerics, Chi-square for categoricals) plus a "
        "dataset-level verdict computed as a configurable share of "
        "drifted columns.",
        "TargetDriftPreset - the same machinery applied specifically to "
        "the target column, useful for detecting label-prevalence shifts.",
        "DataQualityPreset - per-column descriptive statistics (mean, "
        "std, missing fraction, distinct values) used to profile the "
        "reference set in absolute terms.",
        "Reports rendered as both HTML (for human inspection) and JSON "
        "(for machine consumption by the CI gate).",
    ])
    doc.add_page_break()

    # ---- 3. Architecture and Engine ----
    add_heading(doc, "3. Architecture and Evidently Engine", level=1)
    add_para(doc,
        "Figure 1 on the cover page summarises the end-to-end flow. The "
        "system is organised into four horizontal bands. The top band "
        "covers data sources and preparation. The middle band is the "
        "Evidently engine, encapsulated in src/monitoring/drift.py and "
        "exposing three pure functions: build_data_profile, "
        "build_drift_report and extract_drift_summary. The third band is "
        "the CI workflow; three drift-related steps share the same "
        "engine but differ in their inputs and exit-code semantics. The "
        "bottom band lists the deterministic outputs: five stable files "
        "in reports/evidently/, a GitHub Actions artifact and the binary "
        "PASS / FAIL verdict that gates the downstream Docker job.")
    add_heading(doc, "3.1 Engine module - src/monitoring/drift.py",
                level=2)
    add_para(doc,
        "The engine is intentionally thin. build_data_profile takes a "
        "single DataFrame and returns an Evidently Report containing a "
        "DataQualityPreset evaluation. build_drift_report takes a "
        "reference and a current DataFrame and returns a Report combining "
        "DataDriftPreset with TargetDriftPreset (the latter is added only "
        "when the target column is present in both inputs). save_report "
        "writes any Report to disk under stable filenames "
        "(profile.html / drift.html and matching .json files), "
        "overwriting the previous run so the directory contains exactly "
        "five files at any time. extract_drift_summary parses the JSON "
        "form of a drift report and returns the four headline figures the "
        "CI gate uses for its pass/fail decision.")
    add_para(doc,
        "Column types are declared explicitly via Evidently's "
        "ColumnMapping, sourced from the same constants used by "
        "src/data/preprocess.py. This guarantees that the drift detector "
        "and the production preprocessing pipeline always agree on which "
        "features are numeric, categorical and binary, eliminating a "
        "common source of false-positive drift caused by misclassified "
        "columns.")
    next_idx = add_screenshots(doc, SHOTS_ENGINE, start_idx=2)
    doc.add_page_break()


    # ---- 4. Frozen Reference ----
    add_heading(doc, "4. The Frozen Reference Pattern", level=1)
    add_para(doc,
        "A drift comparison is only meaningful relative to a stable "
        "baseline. The pattern adopted here is the frozen reference: at "
        "the moment the production model is shipped, the exact training "
        "set used to fit it is copied to data/reference/reference_train_"
        "<version>.csv and committed to git alongside a JSON metadata "
        "sidecar. From that point onwards, every CI run compares today's "
        "regenerated training data against the same committed snapshot. "
        "The snapshot is treated as immutable; changes are made by "
        "freezing a new version, never by editing in place.")
    add_para(doc,
        "The metadata sidecar (reference_train_<version>.metadata.json) "
        "captures full provenance: semantic version, source path, SHA-256 "
        "of the snapshot, git commit at freeze time, row and column "
        "counts, the inferred per-column schema, the target column name "
        "and the target prevalence. SHA-256 is the audit anchor - any "
        "future divergence from the snapshot can be traced back to a "
        "specific git commit and a specific data hash.")
    add_para(doc, "The freeze is automated by scripts/freeze_reference.py:")
    add_bullets(doc, [
        "Input: any training CSV (defaults to data/processed/train.csv).",
        "Output: snapshot CSV, metadata JSON, both committed to git.",
        "Refuses to overwrite an existing version unless --force is "
        "passed, preventing accidental baseline corruption.",
        "Captures the current git HEAD so the snapshot is traceable to "
        "the model code version that produced it.",
    ])
    add_para(doc,
        "The current baseline used by the CI gate is v1.0.0, with 242 "
        "rows, 14 columns, target prevalence 0.4587 and SHA-256 prefix "
        "0e538464. Bumping the baseline is a one-line change in the CI "
        "workflow once a new version has been frozen and committed.")
    next_idx = add_screenshots(doc, SHOTS_REFERENCE, start_idx=next_idx)
    doc.add_page_break()

    # ---- 5. CI Gating ----
    add_heading(doc, "5. CI Gating - Three Trigger Modes", level=1)
    add_para(doc,
        ".github/workflows/ci.yml hosts three drift-related steps that "
        "share the same Evidently engine but differ in their inputs and "
        "in whether they are allowed to fail the build. The matrix below "
        "summarises which steps run for each trigger.")
    add_bullets(doc, [
        "Push or pull request - only the real gate runs. Reference is "
        "v1.0.0 frozen snapshot; current is the freshly regenerated "
        "train.csv. On a healthy main branch this step always reports "
        "zero drift and exits zero, keeping CI green.",
        "Manual workflow_dispatch with demo_fail_on_drift = false - the "
        "real gate plus an auxiliary demo step that runs Evidently "
        "against synthetically drifted data with --no-fail-on-drift. The "
        "auxiliary step proves the report works against a known-drifted "
        "input but does not block the build.",
        "Manual workflow_dispatch with demo_fail_on_drift = true - the "
        "real gate plus the failure demo step. The failure demo runs the "
        "same drift check WITHOUT --no-fail-on-drift, so a positive "
        "verdict propagates through sys.exit and CI turns red. This is "
        "used to demonstrate the gate's rejection behaviour.",
    ])
    add_para(doc,
        "Two implementation details warrant note. First, the conditional "
        "for the failure step uses truthy evaluation of the input rather "
        "than a literal == true comparison; boolean workflow_dispatch "
        "inputs are sometimes stringified by GitHub, which causes the "
        "literal comparison to silently mis-evaluate. Second, the "
        "Evidently report upload step runs unconditionally (if: always), "
        "so the HTML reports remain downloadable from the run page even "
        "when the build fails. Engineers can therefore inspect what "
        "drifted without having to re-run anything locally.")
    next_idx = add_screenshots(doc, SHOTS_CI, start_idx=next_idx)
    doc.add_page_break()

    # ---- 6. Demonstration ----
    add_heading(doc, "6. Demonstration and Results", level=1)
    add_heading(doc, "6.1 Healthy path - zero drift", level=2)
    add_para(doc,
        "Running the gate against the frozen reference and the freshly "
        "regenerated train.csv produces dataset_drift = False and "
        "share_of_drifted_columns = 0.0000 in 14 of 14 columns. Files are "
        "byte-identical at SHA-256 0e538464 and the script exits zero, "
        "letting CI proceed to the Docker build job. The same outcome "
        "holds locally and in CI because requirements.txt pins all "
        "scientific-Python dependencies, making the prepare.py output "
        "deterministic across platforms.")
    next_idx = add_screenshots(doc, SHOTS_HEALTHY, start_idx=next_idx)
    add_heading(doc, "6.2 Drifted path - gate rejects the build", level=2)
    add_para(doc,
        "scripts/generate_drift_demo.py fabricates a current dataset "
        "designed to trigger every detector: numeric features are shifted "
        "by 1.5 standard deviations, ~30 percent of categorical entries "
        "are flipped to a different valid level, and the target prevalence "
        "is skewed downward. Running run_evidently.py against this dataset "
        "yields dataset_drift = True with share_of_drifted_columns = "
        "0.7143 (10 of 14 columns flagged). The script exits with code 1, "
        "GitHub marks the step failed, and the dependent Docker build job "
        "is skipped. Despite the failure, the artifact upload step still "
        "runs and the HTML reports remain attached to the run page.")
    next_idx = add_screenshots(doc, SHOTS_DRIFTED, start_idx=next_idx)
    add_heading(doc, "6.3 Side-by-side outcomes and visual reports", level=2)
    add_para(doc,
        "The most informative artefact for a stakeholder review is the "
        "pair of CI runs on the same commit - one green from the push "
        "trigger, one red from the manual trigger - with the drift HTML "
        "report attached to the failed run. The HTML walks the reader "
        "from the dataset-level verdict at the top, through a per-column "
        "table with the chosen statistical test and its p-value, down to "
        "individual reference-versus-current distribution overlays. This "
        "is the same evidence path a downstream MLOps engineer would "
        "follow when investigating a real drift incident in production.")
    next_idx = add_screenshots(doc, SHOTS_OUTPUTS, start_idx=next_idx)
    doc.add_page_break()

    # ---- 7. Conclusion ----
    add_heading(doc, "7. Conclusion and Future Work", level=1)
    add_para(doc,
        "Evidently AI has been integrated into the Heart Disease MLOps "
        "pipeline as a CI-gated data-quality and drift detector. The "
        "integration is small (one engine module, three scripts, three "
        "CI steps, seven new pytest cases) and additive (no changes to "
        "the model, container or deployment surface). The frozen "
        "reference pattern gives the gate something stable to compare "
        "against; the workflow_dispatch toggle gives stakeholders a "
        "one-click way to demonstrate both healthy and rejection paths "
        "on the same commit; and the unconditional artifact upload "
        "preserves the evidence trail even on failure.")
    add_para(doc,
        "Three extensions are earmarked for the next iteration:")
    add_bullets(doc, [
        "Live drift endpoint - expose /drift on the Flask app that "
        "compares a rolling window of /predict inputs against the frozen "
        "reference, surfacing share_of_drifted_columns as a Prometheus "
        "gauge and as a Grafana panel.",
        "Scheduled drift check - a separate workflow on a nightly cron "
        "that pulls a real production sample and runs the same gate, "
        "decoupling drift detection from code pushes.",
        "Second DMML tool - layer Great Expectations or Pandera on top "
        "of the same reference snapshot to add schema and value-range "
        "assertions that complement Evidently's distributional checks.",
    ])
    return doc


# ------------------------------ entry point ------------------------------

def main() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    SHOT_DIR.mkdir(parents=True, exist_ok=True)
    arch_path = FIG_DIR / "dmml_architecture.png"
    print(f"[1/2] writing architecture diagram -> {arch_path}")
    draw_architecture(arch_path)

    print(f"[2/2] writing report -> {OUT_DOCX}")
    doc = build_report()
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    size_kb = OUT_DOCX.stat().st_size / 1024
    print(f"        wrote {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
